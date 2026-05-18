#!/usr/bin/env python3
"""Build a formatted .docx paper from a content-spec JSON + a .docx template.

The template supplies every visual decision: named paragraph styles, fonts,
sizes, line spacing, page margins, headers/footers, AND its section/column
layout (e.g. a 1-column title band over a 2-column body).

This script keeps the template's section skeleton intact. It removes the
template's placeholder *content* but preserves every section-break paragraph,
so the original column layout survives. Content blocks are then inserted into
the correct section; a "section_break" block advances to the next section.

Usage:
    python build_paper.py <spec.json>

See SKILL.md for the full spec JSON schema. Summary:

{
  "template": "form.docx",
  "output":   "result.docx",
  "image_base_dir": "figures",          # optional; resolves relative images
  "style_map": {                         # template style names per block type
    "title": "paper title", "subtitle": "paper subtitle",
    "author": "Author", "affiliation": "Affiliation", "abstract": "Abstract",
    "keywords": "Keywords", "heading1": "Heading 1", "heading2": "Heading 2",
    "heading3": "Heading 3", "body": "Body Text", "caption": "figure caption",
    "reference": "references", "equation": "equation"
  },
  "blocks": [
    {"type": "title",         "text": "..."},
    {"type": "section_break"},                       # advance to next section
    {"type": "author",        "text": "Name\\ndept\\norg\\ncity\\nemail"},
    {"type": "heading", "level": 1, "text": "1. Introduction"},
    {"type": "paragraph",     "text": "body text, inline math $a=b$ allowed"},
    {"type": "equation",      "latex": "E = mc^2", "number": "(1)"},
    {"type": "figure",        "image": "fig1.png", "caption": "Fig. 1. ...",
                              "width_cm": 8.5, "caption_position": "below"},
    {"type": "table",         "caption": "Table 1. ...", "caption_position": "above",
                              "header": ["A", "B"], "rows": [["1", "2"]],
                              "table_style": "Table Grid"},
    {"type": "reference",     "text": "[1] Author, Title, Journal, Year."},
    {"type": "pagebreak"}
  ]
}

Newlines (\\n) inside any text become line breaks. Any block may carry an
explicit "style" key to override the style_map.

Requires: python-docx, lxml   (equations also need latex2mathml + mathml2omml)
"""
import os
import re
import sys
import json

# Windows consoles often default to cp949/cp1252; warnings may contain
# em-dashes or non-ASCII paths. Force UTF-8 so printing never crashes.
try:
    sys.stdout.reconfigure(encoding="utf-8")
    sys.stderr.reconfigure(encoding="utf-8")
except (AttributeError, ValueError):
    pass

try:
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.oxml.ns import qn
except ImportError:
    sys.exit("python-docx is required:  pip install python-docx")

from lxml import etree

WARNINGS = []


def warn(msg):
    WARNINGS.append(msg)
    print(f"  [warn] {msg}")


# --------------------------------------------------------------------------
# Equations: LaTeX -> OMML (native Word equations)
# --------------------------------------------------------------------------
_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def latex_to_omml_element(latex):
    """Return an lxml <m:oMath> element for the given LaTeX, or None on failure."""
    try:
        from latex2mathml.converter import convert as latex2mathml
        import mathml2omml
    except ImportError:
        warn("latex2mathml / mathml2omml not installed -- equation kept as text. "
             "Install:  pip install latex2mathml mathml2omml")
        return None
    try:
        mathml = latex2mathml(latex)
        omml = mathml2omml.convert(mathml)
        if "xmlns:m=" not in omml:
            omml = omml.replace("<m:oMath", f'<m:oMath xmlns:m="{_M_NS}"', 1)
        return etree.fromstring(omml)
    except Exception as e:  # noqa: BLE001 - one bad equation must not abort
        warn(f"equation conversion failed for '{latex}': {e}")
        return None


def add_text_with_inline_math(paragraph, text):
    """Add `text` to a paragraph: \\n -> line break, $...$ -> inline OMML."""
    for li, line in enumerate((text or "").split("\n")):
        if li > 0:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        # split on single-dollar spans (ignore escaped \$)
        parts = re.split(r"(?<!\\)\$(.+?)(?<!\\)\$", line)
        for i, part in enumerate(parts):
            if i % 2 == 0:
                if part:
                    paragraph.add_run(part.replace(r"\$", "$"))
            else:
                el = latex_to_omml_element(part)
                if el is not None:
                    paragraph._p.append(el)
                else:
                    paragraph.add_run(f"${part}$")


# --------------------------------------------------------------------------
# Builder -- keeps the template's section skeleton intact
# --------------------------------------------------------------------------
class DocBuilder:
    """Inserts content into a template while preserving its section breaks.

    The template's placeholder content is removed, but every section-break
    paragraph (the carrier of a column/layout change) is kept as an empty
    'anchor'. Content is inserted *before* the current anchor, so it lands in
    the right section. `section_break()` moves on to the next anchor; after
    the last anchor, content goes into the final section.
    """

    def __init__(self, doc):
        self.doc = doc
        self.anchors = self._strip_body()
        self.sec = 0

    def _strip_body(self):
        body = self.doc.element.body
        anchors = []
        for child in list(body):
            tag = child.tag.split("}")[-1]
            if tag == "sectPr":
                continue  # final body section properties -- keep
            if tag == "p":
                pPr = child.find(qn("w:pPr"))
                if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                    # section-break paragraph: keep it (it carries the layout
                    # change) but empty its content and drop its paragraph
                    # style -- otherwise an auto-numbered style (e.g. a
                    # "references" list) would emit a phantom empty entry.
                    for run in child.findall(qn("w:r")):
                        child.remove(run)
                    for hl in child.findall(qn("w:hyperlink")):
                        child.remove(hl)
                    pStyle = pPr.find(qn("w:pStyle"))
                    if pStyle is not None:
                        pPr.remove(pStyle)
                    numPr = pPr.find(qn("w:numPr"))
                    if numPr is not None:
                        pPr.remove(numPr)
                    anchors.append(child)
                    continue
            body.remove(child)
        return anchors

    def _place(self, element):
        """Move an element (already appended at body end) into the current section."""
        if self.sec < len(self.anchors):
            self.anchors[self.sec].addprevious(element)
        # else: leave at end -> belongs to the final section

    def section_break(self):
        self.sec += 1

    def style_exists(self, name):
        try:
            _ = self.doc.styles[name]
            return True
        except KeyError:
            return False

    def paragraph(self, style_name, fallback="Normal"):
        if style_name and self.style_exists(style_name):
            p = self.doc.add_paragraph(style=style_name)
        else:
            if style_name:
                warn(f"style '{style_name}' not in template -- used '{fallback}'")
            p = self.doc.add_paragraph()
            if self.style_exists(fallback):
                p.style = fallback
        self._place(p._p)
        return p

    def table(self, rows, cols):
        t = self.doc.add_table(rows=rows, cols=cols)
        self._place(t._tbl)
        return t


# --------------------------------------------------------------------------
# Block renderers
# --------------------------------------------------------------------------
def render_text_block(b, block, style_name):
    p = b.paragraph(block.get("style", style_name))
    add_text_with_inline_math(p, block.get("text", ""))


def render_heading(b, block, style_map):
    level = int(block.get("level", 1))
    style_name = block.get("style") or style_map.get(f"heading{level}", f"Heading {level}")
    p = b.paragraph(style_name)
    add_text_with_inline_math(p, block.get("text", ""))


def render_equation(b, block, style_map):
    style_name = block.get("style") or style_map.get("equation") or style_map.get("body", "Normal")
    p = b.paragraph(style_name)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    el = latex_to_omml_element(block.get("latex", ""))
    if el is not None:
        p._p.append(el)
    else:
        p.add_run(block.get("latex", ""))
    if block.get("number"):
        p.add_run("\t" + str(block["number"]))  # tab -> flush-right number


def render_caption(b, block, style_map):
    p = b.paragraph(style_map.get("caption", "Caption"))
    add_text_with_inline_math(p, block.get("caption", ""))


def render_figure(b, block, style_map, base_dir):
    cap_pos = block.get("caption_position", "below")
    if cap_pos == "above" and block.get("caption"):
        render_caption(b, block, style_map)

    img = block.get("image", "")
    if img and not os.path.isabs(img):
        img = os.path.join(base_dir, img)
    p = b.paragraph(style_map.get("body", "Normal"))
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if img and os.path.exists(img):
        kwargs = {}
        if block.get("width_cm"):
            kwargs["width"] = Cm(float(block["width_cm"]))
        try:
            run.add_picture(img, **kwargs)
        except Exception as e:  # noqa: BLE001
            warn(f"could not embed image '{img}': {e}")
            run.add_text(f"[IMAGE ERROR: {block.get('image', '')}]")
    else:
        warn(f"image not found: {img}")
        run.add_text(f"[MISSING IMAGE: {block.get('image', '')}]")

    if cap_pos != "above" and block.get("caption"):
        render_caption(b, block, style_map)


def render_table(b, block, style_map):
    cap_pos = block.get("caption_position", "above")
    if cap_pos == "above" and block.get("caption"):
        render_caption(b, block, style_map)

    header = block.get("header", [])
    rows = block.get("rows", [])
    ncols = len(header) if header else (len(rows[0]) if rows else 1)
    table = b.table(0, ncols)

    tstyle = block.get("table_style")
    if tstyle:
        try:
            table.style = tstyle
        except KeyError:
            warn(f"table style '{tstyle}' not in template")

    if header:
        cells = table.add_row().cells
        for i, h in enumerate(header[:ncols]):
            cells[i].text = str(h)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(list(row)[:ncols]):
            cells[i].text = str(val)

    if cap_pos != "above" and block.get("caption"):
        render_caption(b, block, style_map)


def render_pagebreak(b):
    p = b.paragraph("Normal")
    p.add_run().add_break(WD_BREAK.PAGE)


# --------------------------------------------------------------------------
# Main
# --------------------------------------------------------------------------
DEFAULT_STYLE_MAP = {
    "title": "Title", "subtitle": "Subtitle", "author": "Author",
    "affiliation": "Affiliation", "abstract": "Abstract", "keywords": "Keywords",
    "heading1": "Heading 1", "heading2": "Heading 2", "heading3": "Heading 3",
    "body": "Normal", "caption": "Caption", "reference": "References",
    "equation": "Normal",
}


def build(spec_path):
    with open(spec_path, "r", encoding="utf-8") as f:
        spec = json.load(f)

    template = spec["template"]
    output = spec["output"]
    base_dir = spec.get("image_base_dir") or os.path.dirname(os.path.abspath(spec_path))
    style_map = dict(DEFAULT_STYLE_MAP)
    style_map.update(spec.get("style_map", {}))

    if not os.path.exists(template):
        sys.exit(f"template not found: {template}")

    doc = Document(template)
    b = DocBuilder(doc)

    print(f"Template : {template}")
    print(f"Output   : {output}")
    print(f"Sections : {len(b.anchors) + 1}  (template has {len(b.anchors)} section breaks)")
    print(f"Blocks   : {len(spec.get('blocks', []))}")

    for block in spec.get("blocks", []):
        btype = block.get("type")
        if btype == "section_break":
            b.section_break()
        elif btype == "title":
            render_text_block(b, block, style_map["title"])
        elif btype == "subtitle":
            render_text_block(b, block, style_map["subtitle"])
        elif btype == "author":
            render_text_block(b, block, style_map["author"])
        elif btype == "affiliation":
            render_text_block(b, block, style_map["affiliation"])
        elif btype == "abstract":
            render_text_block(b, block, style_map["abstract"])
        elif btype == "keywords":
            render_text_block(b, block, style_map["keywords"])
        elif btype == "heading":
            render_heading(b, block, style_map)
        elif btype == "paragraph":
            render_text_block(b, block, style_map["body"])
        elif btype == "reference":
            render_text_block(b, block, style_map["reference"])
        elif btype == "equation":
            render_equation(b, block, style_map)
        elif btype == "figure":
            render_figure(b, block, style_map, base_dir)
        elif btype == "table":
            render_table(b, block, style_map)
        elif btype == "pagebreak":
            render_pagebreak(b)
        else:
            warn(f"unknown block type: {btype}")

    doc.save(output)
    print(f"\nSaved: {output}")
    if WARNINGS:
        print(f"\n{len(WARNINGS)} warning(s) -- review and note them in the change log.")
    else:
        print("\nNo warnings.")


if __name__ == "__main__":
    if len(sys.argv) != 2:
        sys.exit("Usage: python build_paper.py <spec.json>")
    build(sys.argv[1])
